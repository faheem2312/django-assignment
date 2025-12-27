from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_docx(output_path):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)  

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    def centered(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_spacing(paragraph, space=6):
        paragraph.paragraph_format.space_after = Pt(space)

    def section_row(table, title):
        row = table.add_row().cells
        row[0].merge(row[1])
        p = row[0].paragraphs[0]
        p.add_run(title).bold = True
        add_spacing(p, 8)

    def data_row(table, left, right):
        row = table.add_row().cells

        p1 = row[0].paragraphs[0]
        p1.add_run(left)
        add_spacing(p1)

        p2 = row[1].paragraphs[0]
        p2.add_run(right)
        add_spacing(p2, 10)

    centered("FORM ‘A’", True)
    centered("MEDIATION APPLICATION FORM", True)
    centered("[REFER RULE 3(1)]", True)
    centered("Mumbai District Legal Services Authority")
    centered("City Civil Court, Mumbai")

    doc.add_paragraph("")

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.3)

    section_row(table, "DETAILS OF PARTIES:")

    data_row(table, "Name of Applicant", "{{client_name}}")

    section_row(table, "Address and contact details of Applicant")

    data_row(
        table,
        "Address",
        "REGISTERED ADDRESS:\n{{branch_address}}\n\n"
        "CORRESPONDENCE BRANCH ADDRESS:\n{{branch_address}}"
    )

    data_row(table, "Telephone No.", "{{telephone}}")
    data_row(table, "Mobile No.", "{{mobile}}")
    data_row(table, "Email ID", "info@kslegal.co.in")

    section_row(
        table,
        "Name, Address and Contact details of Opposite Party:"
    )

    data_row(table, "Name", "{{customer_name}}")

    data_row(
        table,
        "Address",
        "REGISTERED ADDRESS:\n____________________\n\n"
        "CORRESPONDENCE BRANCH ADDRESS:\n____________________"
    )

    data_row(table, "Telephone No.", "{{telephone}}")
    data_row(table, "Mobile No.", "{{mobile}}")
    data_row(table, "Email ID", "xyz@gmail.com")

    section_row(table, "DETAILS OF DISPUTE:")

    section_row(
        table,
        "THE COMM. COURTS (PRE-INSTITUTION SETTLEMENT) RULES, 2018"
    )

    section_row(
        table,
        "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015:"
    )

    doc.save(output_path)
